import streamlit as st
import pandas as pd
from datetime import datetime
import os
from io import BytesIO

# ============ KONFIGURASI ============
st.set_page_config(page_title="Catatan Pelanggaran Santri", layout="wide", initial_sidebar_state="collapsed")

# ============ SESSION STATE ============
if "login" not in st.session_state:
    st.session_state.login = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "role" not in st.session_state:
    st.session_state.role = ""
if "show_delete" not in st.session_state:
    st.session_state.show_delete = False
if "confirm_delete" not in st.session_state:
    st.session_state.confirm_delete = False
if "confirm_single_delete" not in st.session_state:
    st.session_state.confirm_single_delete = False
if "single_delete_index" not in st.session_state:
    st.session_state.single_delete_index = None
if "pelanggaran_terpilih" not in st.session_state:
    st.session_state.pelanggaran_terpilih = []

# ============ FUNGSI CRUD ============
FILE_DATA = "data_pelanggaran.csv"

def baca_data():
    try:
        df = pd.read_csv(FILE_DATA)
        df["Tanggal"] = pd.to_datetime(df["Tanggal"])
        return df
    except FileNotFoundError:
        df = pd.DataFrame(columns=["Tanggal", "Nama Santri", "Pelanggaran", "Denda (Rp)", "Status", "Keterangan"])
        df.to_csv(FILE_DATA, index=False)
        return df

def simpan_data(nama, pelanggaran, denda, status, keterangan=""):
    df = baca_data()
    data_baru = pd.DataFrame({
        "Tanggal": [datetime.now()],
        "Nama Santri": [nama],
        "Pelanggaran": [pelanggaran],
        "Denda (Rp)": [denda],
        "Status": [status],
        "Keterangan": [keterangan]
    })
    df = pd.concat([df, data_baru], ignore_index=True)
    df.to_csv(FILE_DATA, index=False)
    return df

def hapus_data(index_list):
    df = baca_data()
    df = df.drop(index_list).reset_index(drop=True)
    df.to_csv(FILE_DATA, index=False)
    return df

# ============ FUNGSI EXPORT ============
def export_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name="Data Pelanggaran", index=False)
        workbook = writer.book
        worksheet = writer.sheets["Data Pelanggaran"]
        
        from openpyxl.styles import numbers
        for row in range(2, len(df) + 2):
            cell = worksheet.cell(row=row, column=5)
            cell.number_format = '#,##0'
        
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
    
    output.seek(0)
    return output

def export_word(df):
    output = BytesIO()
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    title = doc.add_heading('📋 Laporan Data Pelanggaran Santri', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Tanggal Export: {datetime.now().strftime("%d %B %Y %H:%M")}')
    doc.add_paragraph('')
    
    doc.add_heading('📊 Ringkasan', level=1)
    doc.add_paragraph(f'Total Pelanggaran: {len(df)}')
    doc.add_paragraph(f'Total Denda: Rp {df["Denda (Rp)"].sum():,}')
    doc.add_paragraph('')
    
    doc.add_heading('📋 Data Pelanggaran', level=1)
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['Tanggal', 'Nama Santri', 'Pelanggaran', 'Denda (Rp)', 'Status', 'Keterangan']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Tanggal'].strftime('%d/%m/%Y %H:%M') if isinstance(row['Tanggal'], pd.Timestamp) else str(row['Tanggal'])
        row_cells[1].text = str(row['Nama Santri'])
        row_cells[2].text = str(row['Pelanggaran'])
        row_cells[3].text = f"Rp {row['Denda (Rp)']:,}"
        row_cells[4].text = str(row['Status'])
        row_cells[5].text = str(row['Keterangan']) if pd.notna(row['Keterangan']) else ""
    
    doc.save(output)
    output.seek(0)
    return output

# ============ HALAMAN LOGIN ============
def halaman_login():
    st.markdown("""
        <style>
            .stApp { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
            .login-container {
                max-width: 440px;
                margin: auto;
                padding: 40px 35px;
                border-radius: 20px;
                background: rgba(255,255,255,0.95);
                box-shadow: 0 20px 60px rgba(0,0,0,0.3);
                margin-top: 60px;
            }
            .login-title {
                text-align: center;
                margin-bottom: 5px;
                color: #1a73e8;
                font-size: 28px;
                font-weight: 700;
            }
            .login-subtitle {
                text-align: center;
                color: #888;
                margin-bottom: 25px;
                font-size: 14px;
            }
            .login-icon {
                text-align: center;
                font-size: 50px;
                margin-bottom: 5px;
            }
            .divider {
                display: flex;
                align-items: center;
                margin: 20px 0;
                color: #aaa;
                font-size: 13px;
            }
            .divider::before, .divider::after {
                content: "";
                flex: 1;
                border-bottom: 1px solid #ddd;
            }
            .divider::before { margin-right: 15px; }
            .divider::after { margin-left: 15px; }
            .stTextInput > div > div > input {
                border-radius: 10px !important;
                border: 2px solid #e0e0e0 !important;
                padding: 12px 15px !important;
            }
            .stTextInput > div > div > input:focus {
                border-color: #1a73e8 !important;
                box-shadow: 0 0 0 3px rgba(26,115,232,0.2) !important;
            }
            .stButton > button {
                border-radius: 10px !important;
                padding: 12px !important;
                font-weight: 600 !important;
                transition: all 0.3s !important;
            }
            .btn-admin > button {
                background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
                color: white !important;
                border: none !important;
            }
            .btn-admin > button:hover {
                transform: translateY(-2px) !important;
                box-shadow: 0 8px 25px rgba(102, 126, 234, 0.4) !important;
            }
            .btn-tamu > button {
                background: #2ecc71 !important;
                color: white !important;
                border: none !important;
            }
            .btn-tamu > button:hover {
                transform: translateY(-2px) !important;
                box-shadow: 0 8px 25px rgba(46, 204, 113, 0.4) !important;
            }
        </style>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
            <div class="login-container">
                <div class="login-icon">🏫</div>
                <h1 class="login-title">🔐 Login</h1>
                <p class="login-subtitle">Pilih role untuk masuk ke aplikasi</p>
        """, unsafe_allow_html=True)
        
        st.markdown("### 👑 Admin")
        st.caption("Akses penuh: input, edit, hapus, download data")
        
        with st.form("login_admin"):
            username = st.text_input("Username", placeholder="Masukkan username admin", key="login_username")
            password = st.text_input("Password", type="password", placeholder="Masukkan password admin", key="login_password")
            submit_admin = st.form_submit_button("🚀 Login sebagai Admin", use_container_width=True)
            
            if submit_admin:
                if username == "admin" and password == "admin123":
                    st.session_state.login = True
                    st.session_state.username = username
                    st.session_state.role = "admin"
                    st.success("✅ Login sebagai Admin berhasil!")
                    st.rerun()
                else:
                    st.error("❌ Username atau password salah!")
        
        st.markdown('<div class="divider">atau</div>', unsafe_allow_html=True)
        
        st.markdown("### 👤 Pengunjung")
        st.caption("Hanya bisa melihat data (read-only)")
        
        if st.button("👀 Masuk sebagai Pengunjung", use_container_width=True, key="btn_tamu"):
            st.session_state.login = True
            st.session_state.username = "Pengunjung"
            st.session_state.role = "pengunjung"
            st.success("✅ Masuk sebagai Pengunjung!")
            st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)

def cek_akses():
    return st.session_state.role == "admin"

# ============ HALAMAN UTAMA ============
def halaman_utama():
    is_admin = cek_akses()
    
    st.markdown("""
        <style>
            .header {
                display: flex;
                align-items: center;
                justify-content: center;
                background: linear-gradient(135deg, #1a73e8, #0d47a1);
                padding: 15px;
                border-radius: 15px;
                margin-bottom: 20px;
                box-shadow: 0 4px 15px rgba(26,115,232,0.3);
            }
            .header h1 { color: white; margin: 0; font-size: 24px; }
            .header p { color: #e3f2fd; margin: 0; font-size: 14px; }
            .card {
                background: white;
                padding: 20px;
                border-radius: 12px;
                box-shadow: 0 2px 10px rgba(0,0,0,0.05);
                margin-bottom: 15px;
            }
            .card-title {
                font-size: 18px;
                font-weight: 600;
                color: #1a73e8;
                margin-bottom: 10px;
            }
            .badge-admin {
                background: #667eea;
                color: white;
                padding: 4px 12px;
                border-radius: 20px;
                font-size: 12px;
                font-weight: 600;
            }
            .badge-tamu {
                background: #2ecc71;
                color: white;
                padding: 4px 12px;
                border-radius: 20px;
                font-size: 12px;
                font-weight: 600;
            }
        </style>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 3, 1])
    with col1:
        st.write("")
    with col2:
        role_badge = "👑 Admin" if is_admin else "👤 Pengunjung (Read-Only)"
        badge_class = "badge-admin" if is_admin else "badge-tamu"
        st.markdown(f"""
            <div class="header">
                <div style="margin-right:15px; font-size:40px;">🏫</div>
                <div>
                    <h1>📋 Aplikasi Pelanggaran Santri</h1>
                    <p>Selamat datang, {st.session_state.username}! 👋 | <span class="{badge_class}">{role_badge}</span> | 📅 {datetime.now().strftime('%d %B %Y')}</p>
                </div>
            </div>
        """, unsafe_allow_html=True)
    with col3:
        if st.button("🚪 Logout", use_container_width=True, key="btn_logout"):
            st.session_state.login = False
            st.session_state.role = ""
            st.rerun()
    
    st.write("---")
    
    col_input, col_data = st.columns([1, 2])
    
    # ============ INPUT (ADMIN ONLY) ============
    with col_input:
        if is_admin:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-title">📝 Tambah Pelanggaran</div>', unsafe_allow_html=True)
            st.markdown("---")
            
            denda_dict = {"Sholat Jamaah": 20000, "Asmaul Husna": 5000}
            
            nama = st.text_input("👤 Nama Santri", placeholder="Masukkan nama santri", key="nama_input")
            
            pelanggaran_terpilih = st.multiselect(
                "📋 Jenis Pelanggaran (bisa pilih lebih dari 1)",
                ["Sholat Jamaah", "Asmaul Husna"],
                default=st.session_state.pelanggaran_terpilih,
                key="pelanggaran_multiselect"
            )
            st.session_state.pelanggaran_terpilih = pelanggaran_terpilih
            
            if pelanggaran_terpilih:
                total_denda = sum(denda_dict[p] for p in pelanggaran_terpilih)
                rincian = " + ".join([f"{p} (Rp {denda_dict[p]:,})" for p in pelanggaran_terpilih])
                st.success(f"💡 **Total Denda:** {rincian} = **Rp {total_denda:,}**")
            else:
                total_denda = 0
                st.warning("⚠️ Pilih minimal 1 pelanggaran")
            
            status = st.selectbox("📌 Status", ["Belum Dibayar", "Sudah Dibayar"], key="status_input")
            keterangan = st.text_area("📝 Keterangan Tambahan", placeholder="Catatan tambahan...", key="keterangan_input")
            
            if st.button("💾 Simpan", type="primary", use_container_width=True, key="btn_simpan"):
                if nama == "":
                    st.warning("⚠️ Nama santri harus diisi!")
                elif not pelanggaran_terpilih:
                    st.warning("⚠️ Pilih minimal 1 pelanggaran!")
                else:
                    simpan_data(nama, " + ".join(pelanggaran_terpilih), total_denda, status, keterangan)
                    st.success(f"✅ Data berhasil disimpan! Denda: Rp {total_denda:,}")
                    st.rerun()
            
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="card">', unsafe_allow_html=True)
            st.markdown('<div class="card-title">👀 Mode Pengunjung</div>', unsafe_allow_html=True)
            st.markdown("---")
            st.info("ℹ️ Anda login sebagai **Pengunjung**. Hanya bisa melihat data. Login sebagai Admin untuk mengedit.")
            st.markdown('</div>', unsafe_allow_html=True)
    
    # ============ DATA ============
    with col_data:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown('<div class="card-title">📊 Data Pelanggaran</div>', unsafe_allow_html=True)
        st.markdown("---")
        
        df = baca_data()
        
        if df.empty:
            st.info("📭 Belum ada data pelanggaran.")
        else:
            # ============ FILTER ============
            col_filter1, col_filter2, col_filter3, col_filter4 = st.columns(4)
            with col_filter1:
                filter_nama = st.text_input("🔍 Cari Nama Santri", placeholder="Ketik nama...", key="filter_nama")
            with col_filter2:
                filter_tanggal = st.text_input("📅 Cari Tanggal", placeholder="DD/MM/YYYY atau DD/MM...", key="filter_tanggal")
            with col_filter3:
                filter_status = st.selectbox("Filter Status", ["Semua", "Belum Dibayar", "Sudah Dibayar"], key="filter_status")
            with col_filter4:
                pilihan_filter = ["Semua", "Sholat Jamaah", "Asmaul Husna", "Sholat Jamaah + Asmaul Husna"]
                filter_pelanggaran = st.selectbox("Filter Pelanggaran", pilihan_filter, key="filter_pelanggaran")
            
            df_filter = df.copy()
            if filter_nama:
                df_filter = df_filter[df_filter["Nama Santri"].str.contains(filter_nama, case=False)]
            
            if filter_tanggal:
                try:
                    if "/" in filter_tanggal:
                        parts = filter_tanggal.split("/")
                        if len(parts) == 3:
                            filter_date = pd.to_datetime(filter_tanggal, format="%d/%m/%Y", errors="coerce")
                            if pd.notna(filter_date):
                                df_filter = df_filter[df_filter["Tanggal"].dt.date == filter_date.date()]
                        elif len(parts) == 2:
                            day = int(parts[0])
                            month = int(parts[1])
                            df_filter = df_filter[(df_filter["Tanggal"].dt.day == day) & (df_filter["Tanggal"].dt.month == month)]
                    elif filter_tanggal.isdigit() and len(filter_tanggal) == 4:
                        year = int(filter_tanggal)
                        df_filter = df_filter[df_filter["Tanggal"].dt.year == year]
                except:
                    pass
            
            if filter_status != "Semua":
                df_filter = df_filter[df_filter["Status"] == filter_status]
            
            if filter_pelanggaran != "Semua":
                if filter_pelanggaran == "Sholat Jamaah":
                    df_filter = df_filter[df_filter["Pelanggaran"] == "Sholat Jamaah"]
                elif filter_pelanggaran == "Asmaul Husna":
                    df_filter = df_filter[df_filter["Pelanggaran"] == "Asmaul Husna"]
                elif filter_pelanggaran == "Sholat Jamaah + Asmaul Husna":
                    df_filter = df_filter[df_filter["Pelanggaran"] == "Sholat Jamaah + Asmaul Husna"]
            
            if df_filter.empty:
                st.info("📭 Tidak ada data yang sesuai dengan filter")
            else:
                # ============ TAMPILAN DATA ============
                if is_admin:
                    st.info("💡 **Cara edit:** Klik kolom Status, pilih 'Belum Dibayar' atau 'Sudah Dibayar'.")
                    
                    df_edit = df_filter.copy()
                    df_edit["Tanggal"] = df_edit["Tanggal"].dt.strftime("%d/%m/%Y %H:%M")
                    
                    edited_df = st.data_editor(
                        df_edit[["Tanggal", "Nama Santri", "Pelanggaran", "Denda (Rp)", "Status", "Keterangan"]],
                        use_container_width=True,
                        height=350,
                        column_config={
                            "Status": st.column_config.SelectboxColumn(
                                "Status",
                                options=["Belum Dibayar", "Sudah Dibayar"],
                                required=True,
                            ),
                            "Denda (Rp)": st.column_config.NumberColumn(
                                "Denda (Rp)",
                                format="Rp %d",
                            ),
                            "Tanggal": st.column_config.TextColumn("Tanggal", disabled=True),
                            "Nama Santri": st.column_config.TextColumn("Nama Santri", disabled=True),
                            "Pelanggaran": st.column_config.TextColumn("Pelanggaran", disabled=True),
                        },
                        disabled=["Tanggal", "Nama Santri", "Pelanggaran"],
                        num_rows="fixed",
                        key="data_editor"
                    )
                    
                    col_btn1, col_btn2, col_btn3, col_btn4 = st.columns(4)
                    
                    with col_btn1:
                        if st.button("💾 Simpan Perubahan", type="primary", use_container_width=True):
                            for new_idx, row in edited_df.iterrows():
                                original_idx = df_filter.index[new_idx]
                                df.at[original_idx, "Status"] = row["Status"]
                                df.at[original_idx, "Denda (Rp)"] = row["Denda (Rp)"]
                                df.at[original_idx, "Keterangan"] = row["Keterangan"]
                            df.to_csv(FILE_DATA, index=False)
                            st.success("✅ Perubahan berhasil disimpan!")
                            st.rerun()
                    
                    with col_btn2:
                        # ===== DOWNLOAD EXCEL =====
                        if st.button("📥 Download Excel", use_container_width=True):
                            excel_data = export_excel(df)
                            st.download_button(
                                label="💾 Klik untuk Download Excel",
                                data=excel_data,
                                file_name=f"Data_Pelanggaran_{datetime.now().strftime('%d%m%Y')}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key="download_excel"
                            )
                            st.success("✅ File Excel siap didownload!")
                    
                    with col_btn3:
                        # ===== DOWNLOAD WORD =====
                        if st.button("📄 Download Word", use_container_width=True):
                            word_data = export_word(df)
                            st.download_button(
                                label="💾 Klik untuk Download Word",
                                data=word_data,
                                file_name=f"Data_Pelanggaran_{datetime.now().strftime('%d%m%Y')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_word"
                            )
                            st.success("✅ File Word siap didownload!")
                    
                    with col_btn4:
                        if st.button("🗑️ Hapus Massal", use_container_width=True):
                            st.session_state.show_delete = True
                            st.rerun()
                    
                    # ============ HAPUS PER BARIS (EXPANDABLE) ============
                    with st.expander("🗑️ Hapus per Baris (Klik untuk buka/tutup)", expanded=False):
                        st.caption("Klik tombol 🗑️ di samping data yang ingin dihapus")
                        
                        for i in range(len(df_filter)):
                            col_a, col_b, col_c, col_d, col_e, col_f, col_g = st.columns([1, 1.5, 1.5, 1.5, 1.5, 1.5, 0.8])
                            
                            with col_a:
                                st.write(df_filter.iloc[i]["Tanggal"].strftime("%d/%m %H:%M"))
                            with col_b:
                                st.write(df_filter.iloc[i]["Nama Santri"])
                            with col_c:
                                st.write(df_filter.iloc[i]["Pelanggaran"])
                            with col_d:
                                st.write(f"Rp {df_filter.iloc[i]['Denda (Rp)']:,.0f}")
                            with col_e:
                                st.write(df_filter.iloc[i]["Status"])
                            with col_f:
                                st.write(df_filter.iloc[i]["Keterangan"] if pd.notna(df_filter.iloc[i]["Keterangan"]) else "-")
                            with col_g:
                                if st.button("🗑️", key=f"hapus_baris_{i}"):
                                    st.session_state.single_delete_index = df_filter.index[i]
                                    st.session_state.confirm_single_delete = True
                                    st.rerun()
                            st.divider()
                    
                    # ============ KONFIRMASI HAPUS ============
                    if st.session_state.confirm_single_delete:
                        st.error("⚠️ **PERINGATAN!**")
                        st.warning("Apakah Anda yakin ingin menghapus data ini?")
                        st.write(f"**Nama:** {df.loc[st.session_state.single_delete_index, 'Nama Santri']}")
                        st.write(f"**Pelanggaran:** {df.loc[st.session_state.single_delete_index, 'Pelanggaran']}")
                        st.write(f"**Denda:** Rp {df.loc[st.session_state.single_delete_index, 'Denda (Rp)']:,.0f}")
                        
                        col_yes, col_no = st.columns(2)
                        with col_yes:
                            if st.button("✅ Ya, Hapus!", type="primary", use_container_width=True):
                                hapus_data([st.session_state.single_delete_index])
                                st.session_state.confirm_single_delete = False
                                st.session_state.single_delete_index = None
                                st.success("✅ Data berhasil dihapus!")
                                st.rerun()
                        with col_no:
                            if st.button("❌ Tidak", use_container_width=True):
                                st.session_state.confirm_single_delete = False
                                st.session_state.single_delete_index = None
                                st.rerun()
                    
                    # ============ HAPUS MASSAL ============
                    if st.session_state.show_delete:
                        st.warning("⚠️ **Mode Hapus Massal Aktif**")
                        st.write("Centang data yang ingin dihapus, lalu klik tombol 'Hapus yang Dipilih'")
                        
                        df_delete = df_filter.copy()
                        df_delete["Tanggal"] = df_delete["Tanggal"].dt.strftime("%d/%m/%Y %H:%M")
                        
                        selected_indices = []
                        for i in range(len(df_delete)):
                            col_cek, col_tanggal, col_nama, col_pelanggaran, col_denda, col_status = st.columns([0.5, 1.5, 1.5, 1.5, 1.5, 1.5])
                            with col_cek:
                                cek = st.checkbox("", key=f"delete_massal_{i}")
                                if cek:
                                    selected_indices.append(i)
                            with col_tanggal:
                                st.write(df_delete.iloc[i]["Tanggal"])
                            with col_nama:
                                st.write(df_delete.iloc[i]["Nama Santri"])
                            with col_pelanggaran:
                                st.write(df_delete.iloc[i]["Pelanggaran"])
                            with col_denda:
                                st.write(f"Rp {df_delete.iloc[i]['Denda (Rp)']:,.0f}")
                            with col_status:
                                st.write(df_delete.iloc[i]["Status"])
                            st.divider()
                        
                        col_btn1, col_btn2, col_btn3 = st.columns(3)
                        with col_btn1:
                            if st.button("❌ Batal", use_container_width=True):
                                st.session_state.show_delete = False
                                st.session_state.confirm_delete = False
                                st.rerun()
                        with col_btn2:
                            if st.button("🗑️ Hapus yang Dipilih", type="primary", use_container_width=True):
                                if len(selected_indices) == 0:
                                    st.warning("⚠️ Pilih minimal satu data untuk dihapus!")
                                else:
                                    st.session_state.confirm_delete = True
                        with col_btn3:
                            if st.button("☑️ Pilih Semua", use_container_width=True):
                                for i in range(len(df_delete)):
                                    st.session_state[f"delete_massal_{i}"] = True
                                st.rerun()
                        
                        if st.session_state.confirm_delete:
                            st.error("⚠️ **PERINGATAN!**")
                            st.warning(f"Anda akan menghapus **{len(selected_indices)}** data. Tindakan ini tidak dapat dibatalkan!")
                            
                            col_yes, col_no = st.columns(2)
                            with col_yes:
                                if st.button("✅ Ya, Hapus Semua!", type="primary", use_container_width=True):
                                    indices_to_delete = []
                                    for idx in selected_indices:
                                        original_idx = df_filter.index[idx]
                                        indices_to_delete.append(original_idx)
                                    hapus_data(indices_to_delete)
                                    st.session_state.show_delete = False
                                    st.session_state.confirm_delete = False
                                    st.success(f"✅ Berhasil menghapus {len(indices_to_delete)} data!")
                                    st.rerun()
                            with col_no:
                                if st.button("❌ Tidak, Batal", use_container_width=True):
                                    st.session_state.confirm_delete = False
                                    st.rerun()
                
                else:
                    # ===== PENGUNJUNG (READ-ONLY) =====
                    df_tampil = df_filter.copy()
                    df_tampil["Tanggal"] = df_tampil["Tanggal"].dt.strftime("%d/%m/%Y %H:%M")
                    st.dataframe(
                        df_tampil[["Tanggal", "Nama Santri", "Pelanggaran", "Denda (Rp)", "Status", "Keterangan"]],
                        use_container_width=True,
                        height=400
                    )
                
                # ============ RINGKASAN ============
                st.subheader("📈 Ringkasan")
                st.markdown("---")
                
                col_r1, col_r2, col_r3, col_r4 = st.columns(4)
                total_data = len(df)
                total_belum = len(df[df["Status"] == "Belum Dibayar"])
                total_sudah = len(df[df["Status"] == "Sudah Dibayar"])
                total_denda = df["Denda (Rp)"].sum()
                
                col_r1.metric("📝 Total Pelanggaran", total_data)
                col_r2.metric("🔴 Belum Dibayar", total_belum)
                col_r3.metric("🟢 Sudah Dibayar", total_sudah)
                col_r4.metric("💰 Total Denda", f"Rp {total_denda:,.0f}")
                
                # ============ REKAP PER SANTRI (EXPANDABLE) ============
                with st.expander("📊 Rekap per Santri (Klik untuk buka/tutup)", expanded=False):
                    st.markdown("---")
                    
                    cari_santri = st.text_input("🔍 Cari Nama Santri", placeholder="Ketik nama santri...", key="cari_santri_rekap")
                    
                    rekap_orang = df.groupby("Nama Santri").agg({
                        "Pelanggaran": lambda x: ", ".join(x.unique()),
                        "Denda (Rp)": "sum",
                    }).reset_index()
                    
                    rekap_orang["Jumlah Pelanggaran"] = df.groupby("Nama Santri").size().values
                    
                    status_counts = df.groupby("Nama Santri")["Status"].value_counts().unstack(fill_value=0)
                    if "Belum Dibayar" in status_counts.columns:
                        rekap_orang["Belum Dibayar"] = status_counts["Belum Dibayar"].values
                    else:
                        rekap_orang["Belum Dibayar"] = 0
                        
                    if "Sudah Dibayar" in status_counts.columns:
                        rekap_orang["Sudah Dibayar"] = status_counts["Sudah Dibayar"].values
                    else:
                        rekap_orang["Sudah Dibayar"] = 0
                    
                    rekap_orang = rekap_orang.rename(columns={
                        "Nama Santri": "Nama Santri",
                        "Pelanggaran": "Pelanggaran",
                        "Denda (Rp)": "Total Denda",
                        "Belum Dibayar": "🔴 Belum Bayar",
                        "Sudah Dibayar": "🟢 Sudah Bayar"
                    })
                    
                    rekap_orang = rekap_orang.sort_values("Total Denda", ascending=False)
                    
                    if cari_santri:
                        rekap_orang = rekap_orang[rekap_orang["Nama Santri"].str.contains(cari_santri, case=False)]
                    
                    if not rekap_orang.empty:
                        st.dataframe(
                            rekap_orang[["Nama Santri", "Pelanggaran", "Jumlah Pelanggaran", "Total Denda", "🔴 Belum Bayar", "🟢 Sudah Bayar"]],
                            use_container_width=True,
                            column_config={
                                "Total Denda": st.column_config.NumberColumn("Total Denda", format="Rp %d"),
                            }
                        )
                        st.caption(f"📌 Total Santri: {len(rekap_orang)} | Total Denda Keseluruhan: Rp {rekap_orang['Total Denda'].sum():,}")
                    else:
                        st.info("📭 Tidak ada data yang sesuai dengan pencarian")
                
                # ============ TOP 10 (EXPANDABLE) ============
                with st.expander("🏆 Top 10 Santri dengan Denda Terbanyak (Klik untuk buka/tutup)", expanded=False):
                    st.markdown("---")
                    
                    top10 = df.groupby("Nama Santri")["Denda (Rp)"].sum().sort_values(ascending=False).head(10)
                    
                    if not top10.empty:
                        top10_df = top10.reset_index()
                        top10_df.columns = ["Nama Santri", "Total Denda"]
                        top10_df["Peringkat"] = range(1, len(top10_df) + 1)
                        
                        def get_medal(rank):
                            if rank == 1:
                                return "🥇"
                            elif rank == 2:
                                return "🥈"
                            elif rank == 3:
                                return "🥉"
                            else:
                                return f"{rank}."
                        
                        top10_df["Peringkat"] = top10_df["Peringkat"].apply(get_medal)
                        
                        st.dataframe(
                            top10_df[["Peringkat", "Nama Santri", "Total Denda"]],
                            use_container_width=True,
                            column_config={
                                "Total Denda": st.column_config.NumberColumn("Total Denda", format="Rp %d"),
                            },
                            hide_index=True
                        )
                        
                        if len(top10_df) > 0:
                            juara1 = top10_df.iloc[0]
                            st.success(f"🏆 **Juara 1:** {juara1['Nama Santri']} dengan total denda **Rp {juara1['Total Denda']:,.0f}**")
                            
                            if len(top10_df) > 1:
                                juara2 = top10_df.iloc[1]
                                st.info(f"🥈 **Juara 2:** {juara2['Nama Santri']} dengan total denda **Rp {juara2['Total Denda']:,.0f}**")
                            
                            if len(top10_df) > 2:
                                juara3 = top10_df.iloc[2]
                                st.info(f"🥉 **Juara 3:** {juara3['Nama Santri']} dengan total denda **Rp {juara3['Total Denda']:,.0f}**")
                    else:
                        st.info("Belum ada data")
        
        st.markdown('</div>', unsafe_allow_html=True)

# ============ MAIN ============
if st.session_state.login:
    halaman_utama()
else:
    halaman_login()
